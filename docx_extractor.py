"""
DOCX extraction using python-docx.

Extracts headings (from Word styles or font-size heuristics), paragraphs,
tables, images, and lists from .docx files, outputting the same PageResult
format as pdf_extractor.py so html_builder can consume them unchanged.
"""

import io
import logging
import re

from docx import Document
from docx.oxml.ns import qn

from pdf_extractor import TextSpan, ImageInfo, TableData, PageResult

logger = logging.getLogger(__name__)

# Default page dimensions (US Letter in points: 612 × 792)
_DEFAULT_WIDTH = 612.0
_DEFAULT_HEIGHT = 792.0

# Synthetic layout constants (points)
_LEFT_MARGIN = 72.0
_LIST_INDENT = 36.0
_LINE_HEIGHT = 16.0

# Heading style name pattern — matches "Heading 1" through "Heading 9"
_HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)

# T058: font-size thresholds for heading inference (no heading styles)
_INFER_H1_MIN_PT = 16.0
_INFER_H2_MIN_PT = 14.0
_INFER_H3_MIN_PT = 12.0  # must also be bold

# Font sizes that map to html_builder's _heading_level() thresholds:
#   h1 >= 24.0, h2 >= 18.0, h3 >= 14.0+bold, h4 >= 12.0+bold
_HEADING_MAPPED_SIZES: dict[int, float] = {
    1: 26.0,
    2: 20.0,
    3: 15.0,
    4: 13.0,
    5: 13.0,
    6: 13.0,
}


# ---------------------------------------------------------------------------
# Paragraph introspection helpers
# ---------------------------------------------------------------------------

def _para_font_size(para) -> float | None:
    """Return the first explicit font size (pt) from runs, or from the style."""
    for run in para.runs:
        if run.font.size is not None:
            return run.font.size.pt
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size.pt
    return None


def _para_bold(para) -> bool:
    """True when every non-empty run is bold."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    return all(bool(r.bold) for r in runs)


def _para_italic(para) -> bool:
    """True when every non-empty run is italic."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    return all(bool(r.italic) for r in runs)


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

def _heading_level_from_style(para) -> int | None:
    """Return heading level (1-6) from a Word heading style, or None."""
    name = para.style.name if para.style else ""
    m = _HEADING_STYLE_RE.match(name)
    if m:
        return min(int(m.group(1)), 6)
    return None


def _heading_level_from_font(para) -> int | None:
    """Infer heading level from font size (T058 — no heading styles).

    Rules:
      ≥ 16 pt        → h1
      ≥ 14 pt        → h2
      ≥ 12 pt + bold → h3
    """
    size = _para_font_size(para)
    if size is None:
        return None
    if size >= _INFER_H1_MIN_PT:
        return 1
    if size >= _INFER_H2_MIN_PT:
        return 2
    if size >= _INFER_H3_MIN_PT and _para_bold(para):
        return 3
    return None


# ---------------------------------------------------------------------------
# List detection
# ---------------------------------------------------------------------------

def _is_list_item(para) -> bool:
    """Detect bulleted or numbered list paragraphs."""
    # Style-name check ("List Bullet", "List Number", "List Paragraph", …)
    name = para.style.name if para.style else ""
    if name.startswith("List"):
        return True
    # numPr XML element (numbering definition reference)
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    return False


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

_IMG_EXT_MAP: dict[str, str] = {
    "image/png": "png",
    "image/jpeg": "jpeg",
    "image/jpg": "jpeg",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
}


def _extract_images(doc: Document) -> list[ImageInfo]:
    """Extract all embedded images via document relationships."""
    images: list[ImageInfo] = []
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            part = rel.target_part
            img_bytes = part.blob
            ext = _IMG_EXT_MAP.get(part.content_type, "png")
            idx = len(images)
            images.append(ImageInfo(
                page_number=0,
                x0=72.0,
                y0=float(idx * 200),
                x1=300.0,
                y1=float(idx * 200 + 150),
                image_bytes=img_bytes,
                extension=ext,
                xref=idx,
            ))
        except Exception:
            logger.warning("Failed to extract DOCX image relationship")
    return images


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

def _extract_tables(doc: Document) -> list[TableData]:
    """Extract tables — first row treated as header."""
    tables: list[TableData] = []
    for tbl in doc.tables:
        all_rows: list[list[str]] = []
        for row in tbl.rows:
            all_rows.append([cell.text.strip() for cell in row.cells])
        if not all_rows:
            continue
        header = all_rows[0]
        data = all_rows[1:] if len(all_rows) > 1 else []
        tables.append(TableData(
            bbox=(72.0, 0.0, 540.0, 100.0),
            header=header,
            rows=data,
        ))
    return tables


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_docx(docx_data: bytes) -> tuple[list[PageResult], dict]:
    """Extract content from a .docx file.

    Returns
    -------
    pages : list[PageResult]
        Single-element list (one virtual page for the whole document).
    metadata : dict
        Document metadata (title, author, subject).
    """
    doc = Document(io.BytesIO(docx_data))

    # --- metadata ---------------------------------------------------------
    metadata: dict = {}
    try:
        cp = doc.core_properties
        metadata["title"] = cp.title or ""
        metadata["author"] = cp.author or ""
        metadata["subject"] = cp.subject or ""
    except Exception:
        pass

    # --- page dimensions --------------------------------------------------
    width, height = _DEFAULT_WIDTH, _DEFAULT_HEIGHT
    try:
        sec = doc.sections[0]
        if sec.page_width:
            width = sec.page_width.pt
        if sec.page_height:
            height = sec.page_height.pt
    except Exception:
        pass

    # --- text spans from paragraphs ---------------------------------------
    text_spans: list[TextSpan] = []
    y = 50.0  # synthetic vertical cursor

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y += _LINE_HEIGHT * 0.5
            continue

        # Heading? (style first, then font-size heuristic — T058)
        hlevel = _heading_level_from_style(para)
        if hlevel is None:
            hlevel = _heading_level_from_font(para)

        font_size = _para_font_size(para) or 11.0
        is_bold = _para_bold(para)
        is_italic = _para_italic(para)
        is_list = _is_list_item(para)

        if hlevel:
            mapped = _HEADING_MAPPED_SIZES.get(hlevel, 13.0)
            is_bold_heading = True
            # h4-h6 need bold flag for html_builder detection
            if hlevel > 3:
                is_bold_heading = True
            text_spans.append(TextSpan(
                text=text,
                x0=_LEFT_MARGIN,
                y0=y,
                x1=_LEFT_MARGIN + len(text) * 7,
                y1=y + mapped,
                font="Arial",
                size=mapped,
                color=0,
                bold=is_bold_heading,
                italic=is_italic,
            ))
            y += mapped + 8.0

        elif is_list:
            span_text = f"\u2022 {text}"
            x0 = _LEFT_MARGIN + _LIST_INDENT
            text_spans.append(TextSpan(
                text=span_text,
                x0=x0,
                y0=y,
                x1=x0 + len(span_text) * 6,
                y1=y + font_size,
                font="Arial",
                size=font_size,
                color=0,
                bold=is_bold,
                italic=is_italic,
            ))
            y += font_size + 4.0

        else:
            text_spans.append(TextSpan(
                text=text,
                x0=_LEFT_MARGIN,
                y0=y,
                x1=_LEFT_MARGIN + len(text) * 6,
                y1=y + font_size,
                font="Arial",
                size=font_size,
                color=0,
                bold=is_bold,
                italic=is_italic,
            ))
            y += font_size + 4.0

    # --- images & tables --------------------------------------------------
    images = _extract_images(doc)
    tables = _extract_tables(doc)

    page = PageResult(
        page_number=0,
        width=width,
        height=height,
        is_scanned=False,  # DOCX is always digital
        text_spans=text_spans,
        images=images,
        tables=tables,
    )

    return [page], metadata
