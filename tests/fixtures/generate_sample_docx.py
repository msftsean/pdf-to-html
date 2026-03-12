"""
Generate tests/fixtures/sample.docx programmatically (T055).

Creates a Word document with:
- Heading 1, 2, 3 styles
- Data table (3 columns × 4 rows)
- Embedded PNG image
- Bulleted list

Run:  python tests/fixtures/generate_sample_docx.py
"""

import io
import os
import struct
import zlib

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _create_small_png(width: int = 10, height: int = 10) -> bytes:
    """Create a minimal valid RGB PNG (solid red)."""
    raw = b""
    for _ in range(height):
        raw += b"\x00"  # filter byte: None
        for _ in range(width):
            raw += b"\xff\x00\x00"  # red pixel
    compressed = zlib.compress(raw)

    def _chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", compressed)
        + _chunk(b"IEND", b"")
    )


def _add_numpr(para, num_id: int = 1, ilvl: int = 0):
    """Attach a numPr element so the paragraph is a list item."""
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.append(numPr)


def generate_sample_docx() -> Document:
    """Build a sample DOCX with all major content types."""
    doc = Document()
    doc.core_properties.title = "NCDIT Sample Document"
    doc.core_properties.author = "Test Suite"

    # --- Heading 1 --------------------------------------------------------
    doc.add_heading("Department of Information Technology", level=1)
    doc.add_paragraph(
        "This sample document demonstrates DOCX extraction capabilities "
        "for the WCAG-compliant document converter."
    )

    # --- Heading 2 --------------------------------------------------------
    doc.add_heading("Budget Overview", level=2)
    doc.add_paragraph(
        "The following table shows the annual budget allocation."
    )

    # --- Data table -------------------------------------------------------
    table = doc.add_table(rows=4, cols=3)
    for row_idx, (cat, amt, pct) in enumerate([
        ("Category", "Amount", "Percentage"),
        ("Personnel", "$1.5M", "65%"),
        ("Operations", "$0.5M", "22%"),
        ("Technology", "$0.3M", "13%"),
    ]):
        table.cell(row_idx, 0).text = cat
        table.cell(row_idx, 1).text = amt
        table.cell(row_idx, 2).text = pct

    # --- Heading 3 --------------------------------------------------------
    doc.add_heading("Strategic Priorities", level=3)

    # --- Bulleted list ----------------------------------------------------
    for item_text in [
        "Modernize legacy systems",
        "Improve cybersecurity posture",
        "Expand digital services",
        "Train workforce on new technologies",
    ]:
        p = doc.add_paragraph(item_text)
        _add_numpr(p)

    # --- Embedded image ---------------------------------------------------
    png_bytes = _create_small_png()
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(2))

    doc.add_paragraph("End of sample document.")
    return doc


if __name__ == "__main__":
    fixtures_dir = os.path.dirname(os.path.abspath(__file__))
    os.makedirs(fixtures_dir, exist_ok=True)
    out = os.path.join(fixtures_dir, "sample.docx")
    generate_sample_docx().save(out)
    print(f"Generated: {out}")
