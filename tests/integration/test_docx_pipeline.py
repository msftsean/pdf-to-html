"""
Integration test: DOCX → extract → html_builder → WCAG validation (T054).

Verifies the full pipeline from a Word document through extraction
and HTML building to WCAG compliance checking.  Confirms that the
html_builder produces accessible output from DOCX-sourced PageResult
objects with zero critical/serious violations.
"""

import io
import re
import struct
import zlib

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_extractor import extract_docx
from html_builder import build_html
from wcag_validator import validate_html


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_small_png(width: int = 4, height: int = 4) -> bytes:
    """Create a minimal valid RGB PNG (solid red)."""
    raw = b""
    for _ in range(height):
        raw += b"\x00"
        for _ in range(width):
            raw += b"\xff\x00\x00"
    compressed = zlib.compress(raw)

    def _chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", compressed)
        + _chunk(b"IEND", b"")
    )


def _add_numpr(para, num_id: int = 1, ilvl: int = 0):
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestDocxPipeline:
    """Full DOCX pipeline: extract → build_html → wcag_validator."""

    def test_simple_docx_zero_critical_serious(self):
        """Basic DOCX with heading + paragraph should pass WCAG."""
        doc = Document()
        doc.add_heading("Annual Report", level=1)
        doc.add_paragraph("This is the introduction.")

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)
        violations = validate_html(html_out)

        critical_serious = [
            v for v in violations if v.severity in ("critical", "serious")
        ]
        assert critical_serious == [], (
            f"WCAG violations: "
            f"{[(v.rule_id, v.description) for v in critical_serious]}"
        )

    def test_html_has_lang(self):
        """Output HTML must have lang attribute."""
        doc = Document()
        doc.add_paragraph("Hello")
        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)
        assert 'lang="en"' in html_out

    def test_html_has_skip_nav_and_main(self):
        """Output must contain skip-nav link and main landmark."""
        doc = Document()
        doc.add_paragraph("Content")
        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)
        assert 'class="skip-nav"' in html_out
        assert 'id="main-content"' in html_out

    def test_headings_rendered(self):
        """Heading styles appear as <h1>/<h2> in generated HTML."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Body text.")

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)

        assert "<h1>" in html_out
        assert "Title" in html_out
        assert "Section" in html_out

    def test_table_accessible(self):
        """Extracted tables produce <th scope='col'> in HTML."""
        doc = Document()
        doc.add_heading("Data", level=1)
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Name"
        tbl.cell(0, 1).text = "Value"
        tbl.cell(1, 0).text = "X"
        tbl.cell(1, 1).text = "42"

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)

        assert "<table" in html_out
        assert '<th scope="col">' in html_out
        assert "Name" in html_out
        assert "42" in html_out

        # No table-related WCAG violations
        violations = validate_html(html_out)
        table_violations = [
            v for v in violations
            if "table" in v.rule_id.lower() or "th" in v.rule_id.lower()
        ]
        assert table_violations == []

    def test_image_accessible(self):
        """Embedded images have alt text and figure/figcaption."""
        doc = Document()
        doc.add_heading("Photos", level=1)
        doc.add_picture(io.BytesIO(_create_small_png()), width=Inches(1))

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta, embed_images=True)

        assert "<figure>" in html_out
        assert "alt=" in html_out

        violations = validate_html(html_out)
        img_violations = [v for v in violations if v.rule_id == "image-alt"]
        assert img_violations == []

    def test_list_rendered(self):
        """List items produce <ul>/<li> in HTML."""
        doc = Document()
        doc.add_heading("Items", level=1)
        for text in ("Apple", "Banana"):
            p = doc.add_paragraph(text)
            _add_numpr(p)

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)

        assert "<ul>" in html_out
        assert "<li>" in html_out
        assert "Apple" in html_out

    def test_full_document_pipeline(self):
        """Realistic multi-section DOCX with all content types passes WCAG."""
        doc = Document()
        doc.core_properties.title = "NCDIT Report"
        doc.add_heading("Department Overview", level=1)
        doc.add_paragraph("This report covers annual activities.")
        doc.add_heading("Budget", level=2)

        tbl = doc.add_table(rows=3, cols=2)
        tbl.cell(0, 0).text = "Category"
        tbl.cell(0, 1).text = "Amount"
        tbl.cell(1, 0).text = "Personnel"
        tbl.cell(1, 1).text = "$1.5M"
        tbl.cell(2, 0).text = "Operations"
        tbl.cell(2, 1).text = "$0.5M"

        doc.add_heading("Priorities", level=2)
        for item in ("Modernize systems", "Improve security", "Train staff"):
            p = doc.add_paragraph(item)
            _add_numpr(p)

        doc.add_picture(io.BytesIO(_create_small_png()), width=Inches(2))

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)
        violations = validate_html(html_out)

        critical_serious = [
            v for v in violations if v.severity in ("critical", "serious")
        ]
        assert critical_serious == [], (
            f"Full pipeline violations: "
            f"{[(v.rule_id, v.description) for v in critical_serious]}"
        )

        # Spot-check content
        assert "Department Overview" in html_out
        assert "<table" in html_out
        assert "<ul>" in html_out
        assert "<figure>" in html_out

    def test_no_heading_styles_uses_font_inference(self):
        """T058: DOCX with manual formatting passes WCAG too."""
        doc = Document()
        para1 = doc.add_paragraph()
        run1 = para1.add_run("Big Title")
        run1.font.size = Pt(18)

        para2 = doc.add_paragraph()
        run2 = para2.add_run("Medium Section")
        run2.font.size = Pt(15)

        para3 = doc.add_paragraph()
        run3 = para3.add_run("Body content here.")
        run3.font.size = Pt(11)

        pages, meta = extract_docx(_docx_bytes(doc))
        html_out, _ = build_html(pages, {}, meta)

        assert "Big Title" in html_out
        assert "Medium Section" in html_out

        # At least one heading tag should have been generated
        headings = re.findall(r"<h[1-6]>", html_out)
        assert len(headings) >= 1

        violations = validate_html(html_out)
        critical_serious = [
            v for v in violations if v.severity in ("critical", "serious")
        ]
        assert critical_serious == []

    def test_sample_docx_fixture_roundtrip(self):
        """If sample.docx exists, verify the full roundtrip."""
        import os
        fixture = os.path.join(
            os.path.dirname(__file__), "..", "fixtures", "sample.docx"
        )
        if not os.path.exists(fixture):
            pytest.skip("sample.docx fixture not generated yet")

        with open(fixture, "rb") as f:
            data = f.read()

        pages, meta = extract_docx(data)
        html_out, _ = build_html(pages, {}, meta)
        violations = validate_html(html_out)

        critical_serious = [
            v for v in violations if v.severity in ("critical", "serious")
        ]
        assert critical_serious == [], (
            f"sample.docx roundtrip violations: "
            f"{[(v.rule_id, v.description) for v in critical_serious]}"
        )
