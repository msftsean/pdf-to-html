"""
Unit tests for docx_extractor — T053

Tests the DOCX extraction pipeline that converts Word documents
to PageResult format compatible with html_builder.

Validates:
  - Heading extraction from Word styles (Heading 1-3)
  - Table extraction
  - Image extraction
  - List extraction (bulleted and numbered)
  - Heading inference from font size (T058)
"""

import io
import struct
import zlib

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_extractor import extract_docx
from pdf_extractor import TextSpan, ImageInfo, TableData, PageResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _docx_bytes(doc: Document) -> bytes:
    """Serialize a python-docx Document to bytes."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_small_png(width: int = 4, height: int = 4) -> bytes:
    """Create a minimal valid RGB PNG (solid red)."""
    raw = b""
    for _ in range(height):
        raw += b"\x00"  # filter byte: None
        for _ in range(width):
            raw += b"\xff\x00\x00"  # red pixel
    compressed = zlib.compress(raw)

    def _chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", compressed)
        + _chunk(b"IEND", b"")
    )


def _add_numpr(para, num_id: int = 1, ilvl: int = 0):
    """Attach a numPr element so the paragraph becomes a list item."""
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Tests: Heading extraction from Word styles
# ---------------------------------------------------------------------------

class TestHeadingExtraction:
    """Heading extraction from Word heading styles (Heading 1-3)."""

    def test_heading1_extracted(self):
        """Heading 1 style maps to a span with size ≥ 24 (h1 in html_builder)."""
        doc = Document()
        doc.add_heading("Main Title", level=1)
        doc.add_paragraph("Body text.")
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages) == 1
        page = pages[0]
        assert page.is_scanned is False

        h1_spans = [s for s in page.text_spans if s.size >= 24.0]
        assert len(h1_spans) >= 1
        assert "Main Title" in h1_spans[0].text

    def test_heading2_extracted(self):
        """Heading 2 style maps to a span with 18 ≤ size < 24 (h2)."""
        doc = Document()
        doc.add_heading("Section Title", level=2)
        pages, _ = extract_docx(_docx_bytes(doc))

        h2_spans = [s for s in pages[0].text_spans if 18.0 <= s.size < 24.0]
        assert len(h2_spans) >= 1
        assert "Section Title" in h2_spans[0].text

    def test_heading3_extracted(self):
        """Heading 3 style maps to a span with 14 ≤ size < 18 and bold."""
        doc = Document()
        doc.add_heading("Subsection", level=3)
        pages, _ = extract_docx(_docx_bytes(doc))

        h3_spans = [
            s for s in pages[0].text_spans
            if 14.0 <= s.size < 18.0 and s.bold
        ]
        assert len(h3_spans) >= 1
        assert "Subsection" in h3_spans[0].text

    def test_multiple_heading_levels_ordered(self):
        """Headings at levels 1-3 appear in correct y-order."""
        doc = Document()
        doc.add_heading("H1 Title", level=1)
        doc.add_heading("H2 Section", level=2)
        doc.add_heading("H3 Detail", level=3)
        doc.add_paragraph("Normal text.")
        pages, _ = extract_docx(_docx_bytes(doc))

        headings = [s for s in pages[0].text_spans if s.size >= 14.0 and s.bold]
        assert len(headings) >= 3
        assert headings[0].y0 < headings[1].y0 < headings[2].y0

    def test_heading_spans_are_bold(self):
        """All heading spans must have bold=True."""
        doc = Document()
        doc.add_heading("Bold Title", level=1)
        pages, _ = extract_docx(_docx_bytes(doc))

        heading = [s for s in pages[0].text_spans if "Bold Title" in s.text][0]
        assert heading.bold is True


# ---------------------------------------------------------------------------
# Tests: Table extraction
# ---------------------------------------------------------------------------

class TestTableExtraction:
    """Table extraction from Word tables."""

    def test_basic_table(self):
        """3×2 table: first row = header, remaining = data."""
        doc = Document()
        tbl = doc.add_table(rows=3, cols=2)
        tbl.cell(0, 0).text = "Name"
        tbl.cell(0, 1).text = "Value"
        tbl.cell(1, 0).text = "Alpha"
        tbl.cell(1, 1).text = "100"
        tbl.cell(2, 0).text = "Beta"
        tbl.cell(2, 1).text = "200"
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages[0].tables) == 1
        td = pages[0].tables[0]
        assert isinstance(td, TableData)
        assert td.header == ["Name", "Value"]
        assert len(td.rows) == 2
        assert td.rows[0] == ["Alpha", "100"]
        assert td.rows[1] == ["Beta", "200"]

    def test_single_row_header_only(self):
        """A table with only one row yields header with no data rows."""
        doc = Document()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.cell(0, 0).text = "Col1"
        tbl.cell(0, 1).text = "Col2"
        pages, _ = extract_docx(_docx_bytes(doc))

        td = pages[0].tables[0]
        assert td.header == ["Col1", "Col2"]
        assert td.rows == []

    def test_empty_cells_preserved(self):
        """Empty table cells are preserved as empty strings."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Header"
        tbl.cell(0, 1).text = ""
        tbl.cell(1, 0).text = ""
        tbl.cell(1, 1).text = "Data"
        pages, _ = extract_docx(_docx_bytes(doc))

        td = pages[0].tables[0]
        assert td.header == ["Header", ""]
        assert td.rows[0] == ["", "Data"]

    def test_multiple_tables(self):
        """Multiple tables in a document are all extracted."""
        doc = Document()
        doc.add_table(rows=2, cols=2)
        doc.add_table(rows=3, cols=3)
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages[0].tables) == 2


# ---------------------------------------------------------------------------
# Tests: Image extraction
# ---------------------------------------------------------------------------

class TestImageExtraction:
    """Image extraction from embedded DOCX images."""

    def test_embedded_png(self):
        """An inline PNG image is extracted as an ImageInfo."""
        doc = Document()
        doc.add_paragraph("Before image.")
        png_bytes = _create_small_png()
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(1))
        doc.add_paragraph("After image.")
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages[0].images) >= 1
        img = pages[0].images[0]
        assert isinstance(img, ImageInfo)
        assert img.extension == "png"
        assert len(img.image_bytes) > 0
        assert img.page_number == 0

    def test_no_images_gives_empty_list(self):
        """A text-only document has an empty images list."""
        doc = Document()
        doc.add_paragraph("No images here.")
        pages, _ = extract_docx(_docx_bytes(doc))

        assert pages[0].images == []


# ---------------------------------------------------------------------------
# Tests: List extraction
# ---------------------------------------------------------------------------

class TestListExtraction:
    """Bulleted and numbered list extraction."""

    def test_list_via_numpr(self):
        """Paragraphs with numPr XML are detected as list items."""
        doc = Document()
        doc.add_paragraph("Intro text.")
        p1 = doc.add_paragraph("First item")
        _add_numpr(p1)
        p2 = doc.add_paragraph("Second item")
        _add_numpr(p2)
        pages, _ = extract_docx(_docx_bytes(doc))

        list_spans = [s for s in pages[0].text_spans if "\u2022" in s.text]
        assert len(list_spans) == 2
        assert "First item" in list_spans[0].text
        assert "Second item" in list_spans[1].text

    def test_list_bullet_style(self):
        """Paragraphs with 'List Bullet' style are detected."""
        doc = Document()
        doc.add_paragraph("Item A", style="List Bullet")
        doc.add_paragraph("Item B", style="List Bullet")
        pages, _ = extract_docx(_docx_bytes(doc))

        list_spans = [s for s in pages[0].text_spans if "\u2022" in s.text]
        assert len(list_spans) == 2

    def test_list_items_indented_vs_body(self):
        """List items have a larger x0 than normal paragraphs."""
        doc = Document()
        doc.add_paragraph("Normal paragraph.")
        p = doc.add_paragraph("List item")
        _add_numpr(p)
        pages, _ = extract_docx(_docx_bytes(doc))

        normal = [s for s in pages[0].text_spans if "Normal" in s.text][0]
        item = [s for s in pages[0].text_spans if "\u2022" in s.text][0]
        assert item.x0 > normal.x0


# ---------------------------------------------------------------------------
# Tests: T058 — Heading inference from font size
# ---------------------------------------------------------------------------

class TestHeadingInference:
    """Heading detection from font size when no heading styles are used."""

    def test_16pt_inferred_as_h1(self):
        """Font ≥ 16pt → h1 (mapped size ≥ 24.0)."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Big Title")
        run.font.size = Pt(18)
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Big Title" in s.text][0]
        assert span.size >= 24.0
        assert span.bold is True

    def test_14pt_inferred_as_h2(self):
        """Font ≥ 14pt but < 16pt → h2 (mapped size ≥ 18.0)."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Medium Title")
        run.font.size = Pt(15)
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Medium Title" in s.text][0]
        assert 18.0 <= span.size < 24.0

    def test_12pt_bold_inferred_as_h3(self):
        """Font ≥ 12pt + bold → h3 (mapped size ≥ 14.0)."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold Section")
        run.font.size = Pt(12)
        run.bold = True
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Bold Section" in s.text][0]
        assert 14.0 <= span.size < 18.0
        assert span.bold is True

    def test_12pt_nonbold_stays_paragraph(self):
        """Font 12pt without bold should NOT be inferred as heading."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Regular text")
        run.font.size = Pt(12)
        run.bold = False
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Regular text" in s.text][0]
        assert span.size == 12.0
        assert span.bold is False

    def test_11pt_stays_paragraph(self):
        """Font 11pt should remain a normal paragraph."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Normal body text")
        run.font.size = Pt(11)
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Normal body text" in s.text][0]
        assert span.size == 11.0

    def test_style_takes_precedence_over_font(self):
        """When a heading style IS present, font-size inference is skipped."""
        doc = Document()
        doc.add_heading("Styled Heading", level=2)
        pages, _ = extract_docx(_docx_bytes(doc))

        span = [s for s in pages[0].text_spans if "Styled Heading" in s.text][0]
        # Should use the h2 mapped size (20.0), not font-size heuristic
        assert span.size == 20.0


# ---------------------------------------------------------------------------
# Tests: Metadata and PageResult structure
# ---------------------------------------------------------------------------

class TestDocxMetadata:
    """Metadata extraction and PageResult structure."""

    def test_title_metadata(self):
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.add_paragraph("Content.")
        _, meta = extract_docx(_docx_bytes(doc))
        assert meta.get("title") == "Test Document"

    def test_author_metadata(self):
        doc = Document()
        doc.core_properties.author = "Jane Doe"
        doc.add_paragraph("Content.")
        _, meta = extract_docx(_docx_bytes(doc))
        assert meta.get("author") == "Jane Doe"

    def test_page_result_structure(self):
        """PageResult has correct type and fields."""
        doc = Document()
        doc.add_paragraph("Hello.")
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages) == 1
        page = pages[0]
        assert isinstance(page, PageResult)
        assert page.page_number == 0
        assert page.width > 0
        assert page.height > 0
        assert page.is_scanned is False

    def test_empty_document(self):
        """An empty DOCX returns an empty PageResult."""
        doc = Document()
        pages, _ = extract_docx(_docx_bytes(doc))

        assert len(pages) == 1
        assert pages[0].text_spans == []
        assert pages[0].tables == []

    def test_is_scanned_always_false(self):
        """DOCX files are always digital — is_scanned must be False."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Text")
        pages, _ = extract_docx(_docx_bytes(doc))
        assert pages[0].is_scanned is False
